#!/usr/bin/env python3
"""
MongoDB Health Check & Design Review CLI

Functions:
- Connect to MongoDB using URI
- Enumerate databases/collections with stats
- Sample documents to infer schema and detect patterns:
  - Field presence and type distribution
  - Array length distribution; flag potential unbounded arrays
  - Document size (BSON) distribution; flag near/over 16MB
  - Nesting depth estimation
- Index inventory with sizes; report largest index per collection
- opcounters-based read/write ratio (if permitted)
- Recommend connection string tweaks based on observed workload
- Emit JSON and Markdown reports

Notes:
- Read-only operations only (stats, sample, aggregate)
- Graceful degradation when commands are not permitted
"""

import argparse
import collections
import datetime
import json
import math
import os
import random
import sys
from typing import Any, Dict, List, Optional, Tuple

from pymongo import MongoClient
from pymongo.errors import PyMongoError
from bson import BSON, json_util
from bson.objectid import ObjectId
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="MongoDB health check and design review")
    parser.add_argument("--uri", required=False, default="", help="MongoDB connection URI (overrides env)")
    parser.add_argument("--env-file", default="", help="Path to .env file to load (optional)")
    parser.add_argument("--dbs", default="", help="Comma-separated database names to include")
    parser.add_argument("--collections", default="", help="Comma-separated collection names to include (applies to provided dbs)")
    parser.add_argument("--sample-size", type=int, default=200, help="Per-collection sample size for schema inference")
    parser.add_argument("--max-docs-per-coll", type=int, default=5000, help="Max docs to scan per collection to avoid load")
    parser.add_argument("--output-dir", default="report", help="Directory to write reports to")
    parser.add_argument("--timeout-ms", type=int, default=10000, help="MongoClient serverSelectionTimeoutMS")
    parser.add_argument("--seed", type=int, default=42, help="Random seed for sampling")
    parser.add_argument("--doc-id", default="", help="Filter a specific document by _id (ObjectId or string)")
    parser.add_argument("--filter", default="", help='JSON query filter applied to sampling, e.g. {"status":"A"}')
    return parser.parse_args()


def get_target_sets(client: MongoClient, include_dbs: List[str], include_colls: List[str]) -> Dict[str, List[str]]:
    target: Dict[str, List[str]] = {}
    try:
        db_names = client.list_database_names()
    except Exception:
        db_names = []
    # exclude system databases by default
    system_dbs = {"admin", "local", "config"}
    for db_name in db_names:
        if db_name in system_dbs:
            continue
        if include_dbs and db_name not in include_dbs:
            continue
        try:
            coll_names = client[db_name].list_collection_names()
        except Exception:
            coll_names = []
        # exclude system collections
        coll_names = [c for c in coll_names if not c.startswith("system.")]
        if include_colls:
            coll_names = [c for c in coll_names if c in include_colls]
        if coll_names:
            target[db_name] = coll_names
    return target


def bson_size(document: Dict[str, Any]) -> int:
    try:
        return len(BSON.encode(document))
    except Exception:
        return -1


def estimate_depth(value: Any, current: int = 0) -> int:
    if isinstance(value, dict):
        if not value:
            return current + 1
        return max(estimate_depth(v, current + 1) for v in value.values())
    if isinstance(value, list):
        if not value:
            return current + 1
        return max(estimate_depth(v, current + 1) for v in value)
    return current + 1


def type_name(value: Any) -> str:
    if value is None:
        return "null"
    if isinstance(value, bool):
        return "bool"
    if isinstance(value, int):
        return "int"
    if isinstance(value, float):
        return "double"
    if isinstance(value, str):
        return "string"
    if isinstance(value, list):
        return "array"
    if isinstance(value, dict):
        return "object"
    return type(value).__name__


def flatten_paths(doc: Dict[str, Any], prefix: str = "") -> List[Tuple[str, Any]]:
    items: List[Tuple[str, Any]] = []
    for k, v in doc.items():
        path = f"{prefix}.{k}" if prefix else k
        if isinstance(v, dict):
            items.extend(flatten_paths(v, path))
        else:
            items.append((path, v))
    return items


def sample_documents(coll, sample_size: int, max_scan: int, rng: random.Random, query_filter: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
    docs: List[Dict[str, Any]] = []
    try:
        # If explicit filter is provided, honor it and avoid $sample to keep semantics
        if query_filter is not None:
            cursor = coll.find(query_filter, projection=None, no_cursor_timeout=False, limit=min(sample_size, max_scan))
            return list(cursor)
        # Attempt $sample first (requires collection not empty and not huge for performance)
        size = min(sample_size, max_scan)
        if size > 0:
            pipeline = [{"$sample": {"size": size}}]
            docs = list(coll.aggregate(pipeline, allowDiskUse=False))
    except Exception:
        docs = []
    if docs:
        return docs
    try:
        cursor = coll.find({}, projection=None, no_cursor_timeout=False, limit=max_scan)
        pool = list(cursor)
        if len(pool) <= sample_size:
            return pool
        idxs = rng.sample(range(len(pool)), sample_size)
        return [pool[i] for i in idxs]
    except Exception:
        return []


def analyze_collection(db_name: str, coll_name: str, coll, sample_size: int, max_scan: int, rng: random.Random, query_filter: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    result: Dict[str, Any] = {
        "db": db_name,
        "collection": coll_name,
        "counts": {},
        "collStats": {},
        "indexSizes": {},
        "largestIndex": None,
        "schema": {
            "fieldPresence": {},
            "fieldTypes": {},
            "arrayStats": {},
            "docSizeBytes": {"max": 0, "p95": 0, "avg": 0},
            "nestingDepth": {"max": 0, "p95": 0, "avg": 0},
            "polymorphism": {},
            "cardinality": {},
        },
        "flags": {
            "unboundedArrays": [],
            "largeDocuments": False,
            "extremeNesting": False,
        },
        "recommendations": [],
        "indexInsights": {
            "highCardinalityCandidates": [],
            "lowCardinalityWarnings": [],
        },
        "notes": [],
    }
    # collStats
    try:
        stats = coll.database.command({"collStats": coll_name})
        result["collStats"] = {
            "count": stats.get("count"),
            "size": stats.get("size"),
            "avgObjSize": stats.get("avgObjSize"),
            "storageSize": stats.get("storageSize"),
            "totalIndexSize": stats.get("totalIndexSize"),
            "indexSizes": stats.get("indexSizes", {}),
        }
        result["indexSizes"] = stats.get("indexSizes", {})
        if result["indexSizes"]:
            largest = max(result["indexSizes"].items(), key=lambda kv: kv[1])
            result["largestIndex"] = {"name": largest[0], "size": largest[1]}
    except Exception as exc:
        result["notes"].append(f"collStats failed: {exc}")

    # counts
    try:
        result["counts"]["documents"] = coll.estimated_document_count()
    except Exception as exc:
        result["notes"].append(f"estimated_document_count failed: {exc}")

    # sample
    samples = sample_documents(coll, sample_size, max_scan, rng, query_filter=query_filter)
    if not samples:
        result["notes"].append("No samples collected; skipping schema analysis")
        return result

    # schema analysis
    field_presence: Dict[str, int] = collections.Counter()
    field_types: Dict[str, collections.Counter] = collections.defaultdict(collections.Counter)
    array_stats_len: Dict[str, List[int]] = collections.defaultdict(list)
    # approximate distinct tracking per field (scalar only), with cap to control memory
    value_sets: Dict[str, set] = collections.defaultdict(set)
    VALUE_CAP = 5000
    doc_sizes: List[int] = []
    depth_vals: List[int] = []

    for doc in samples:
        size_b = bson_size(doc)
        if size_b > 0:
            doc_sizes.append(size_b)
        depth_vals.append(estimate_depth(doc))
        for path, value in flatten_paths(doc):
            field_presence[path] += 1
            tname = type_name(value)
            field_types[path][tname] += 1
            if isinstance(value, list):
                array_stats_len[path].append(len(value))
            # track scalar values for cardinality estimation
            if not isinstance(value, (list, dict)):
                if len(value_sets[path]) < VALUE_CAP:
                    try:
                        # convert to a hashable representation
                        key = json.dumps(value, default=str, ensure_ascii=False)
                    except Exception:
                        key = str(value)
                    value_sets[path].add(key)

    def percentile(data: List[int], p: float) -> int:
        if not data:
            return 0
        s = sorted(data)
        idx = min(len(s) - 1, max(0, int(math.ceil(p * len(s)) - 1)))
        return int(s[idx])

    def avg(data: List[int]) -> float:
        return float(sum(data) / len(data)) if data else 0.0

    schema = result["schema"]
    schema["fieldPresence"] = {k: int(v) for k, v in field_presence.items()}
    schema["fieldTypes"] = {k: dict(v) for k, v in field_types.items()}
    schema["docSizeBytes"] = {
        "max": max(doc_sizes) if doc_sizes else 0,
        "p95": percentile(doc_sizes, 0.95),
        "avg": avg(doc_sizes),
    }
    schema["nestingDepth"] = {
        "max": max(depth_vals) if depth_vals else 0,
        "p95": percentile(depth_vals, 0.95),
        "avg": avg(depth_vals),
    }

    array_stats: Dict[str, Dict[str, Any]] = {}
    unbounded_flags: List[str] = []
    for path, lens in array_stats_len.items():
        stats = {
            "max": max(lens) if lens else 0,
            "p95": percentile(lens, 0.95),
            "avg": avg(lens),
            "samples": len(lens),
        }
        array_stats[path] = stats
        # heuristic threshold
        if stats["max"] >= 1000 or stats["p95"] >= 500:
            unbounded_flags.append(path)
    schema["arrayStats"] = array_stats

    # polymorphism detection
    polymorphism: Dict[str, List[Tuple[str, int]]] = {}
    for path, counter in field_types.items():
        if len(counter) > 1:
            polymorphism[path] = sorted(counter.items(), key=lambda kv: (-kv[1], kv[0]))
    schema["polymorphism"] = {k: [[t, int(c)] for t, c in v] for k, v in polymorphism.items()}

    # cardinality estimation (based on samples)
    cardinality_info: Dict[str, Dict[str, Any]] = {}
    num_samples = max(1, len(samples))
    for path, present_cnt in field_presence.items():
        distinct_cnt = len(value_sets.get(path, set()))
        presence_ratio = float(present_cnt) / float(num_samples)
        ratio = float(distinct_cnt) / float(present_cnt) if present_cnt > 0 else 0.0
        cardinality_info[path] = {
            "present": int(present_cnt),
            "approxDistinct": int(distinct_cnt),
            "distinctPerPresent": ratio,
            "presenceRatio": presence_ratio,
        }
    schema["cardinality"] = cardinality_info

    # flags
    result["flags"]["unboundedArrays"] = unbounded_flags
    if schema["docSizeBytes"]["max"] >= 12 * 1024 * 1024:
        result["flags"]["largeDocuments"] = True
    if schema["nestingDepth"]["max"] > 10:
        result["flags"]["extremeNesting"] = True

    # build recommendations and index insights
    recs: List[str] = []
    idx_high: List[str] = []
    idx_low_warn: List[str] = []

    # schema-related
    if result["flags"]["largeDocuments"]:
        recs.append("Document size approaches/exceeds limits. Consider Split/Reference/Bucket patterns")
    if result["flags"]["extremeNesting"]:
        recs.append("Excessive nesting depth. Consider flattening or embedding only hot fields")
    if result["flags"]["unboundedArrays"]:
        recs.append("Potential unbounded arrays detected. Consider Bucket/Subset/Outlier patterns")
    if polymorphism:
        recs.append("Field type polymorphism detected. Normalize field types or split by schema variant")

    # index-related heuristics based on cardinality
    for path, info in cardinality_info.items():
        present = info["present"]
        presence_ratio = info["presenceRatio"]
        distinct_cnt = info["approxDistinct"]
        distinct_ratio = info["distinctPerPresent"]
        ftypes = field_types.get(path, {})
        # skip arrays/objects as leading candidates, but we still warn
        if "array" in ftypes or "object" in ftypes:
            continue
        if present >= max(30, int(0.3 * num_samples)):
            # high-cardinality suggestion
            if distinct_cnt >= 30 and distinct_ratio >= 0.5:
                idx_high.append(path)
            # low-cardinality warning (avoid as leading index key)
            if distinct_cnt <= 5 and presence_ratio >= 0.5:
                idx_low_warn.append(path)
    if idx_high:
        recs.append(
            "High-cardinality fields frequently present: consider indexing if used as filters -> " + ", ".join(sorted(idx_high))
        )
    if idx_low_warn:
        recs.append(
            "Low-cardinality fields with high presence: avoid leading index position -> " + ", ".join(sorted(idx_low_warn))
        )

    # multikey caution
    for path, stats in array_stats.items():
        if stats.get("max", 0) >= 1000 or stats.get("p95", 0) >= 500:
            recs.append(f"Caution indexing large arrays (multikey explosion): {path}")

    # ESR guidance (conceptual)
    recs.append("Validate index effectiveness via $indexStats and query plans; prefer selective leading keys")

    result["recommendations"] = recs
    result["indexInsights"]["highCardinalityCandidates"] = sorted(idx_high)
    result["indexInsights"]["lowCardinalityWarnings"] = sorted(idx_low_warn)

    return result


def get_server_info(client: MongoClient) -> Dict[str, Any]:
    info: Dict[str, Any] = {"buildInfo": {}, "isMaster": {}, "serverStatus": {}}
    db = client.get_database("admin")
    try:
        info["buildInfo"] = db.command({"buildInfo": 1})
    except Exception:
        pass
    try:
        # hello supersedes isMaster in newer versions; fall back gracefully
        try:
            info["isMaster"] = db.command({"hello": 1})
        except Exception:
            info["isMaster"] = db.command({"isMaster": 1})
    except Exception:
        pass
    try:
        info["serverStatus"] = db.command({"serverStatus": 1})
    except Exception:
        info["serverStatus"] = {}
    return info


def recommend_uri_tweaks(uri: str, server_info: Dict[str, Any]) -> List[str]:
    recs: List[str] = []
    ss = server_info.get("serverStatus", {})
    opc = ss.get("opcounters", {})
    reads = float(opc.get("query", 0) + opc.get("getmore", 0))
    writes = float(opc.get("insert", 0) + opc.get("update", 0) + opc.get("delete", 0))
    # read/write ratio
    ratio = None
    if reads or writes:
        ratio = reads / max(1.0, writes)
    if ratio is not None:
        if ratio >= 5.0:
            recs.append("Workload appears read-heavy; consider readPreference=secondaryPreferred if acceptable")
        elif ratio <= 0.2:
            recs.append("Workload appears write-heavy; prefer primary reads and ensure retryWrites=true")
    # retryWrites
    recs.append("Ensure retryWrites=true (default on Atlas) for idempotent writes")
    # writeConcern
    recs.append("Set writeConcern to match durability needs (e.g., w=majority for strong durability)")
    # pool sizing
    recs.append("Tune maxPoolSize based on client concurrency and server capacity")
    return recs


def write_reports(output_dir: str, report: Dict[str, Any]) -> None:
    os.makedirs(output_dir, exist_ok=True)
    json_path = os.path.join(output_dir, "report.json")
    md_path = os.path.join(output_dir, "report.md")
    with open(json_path, "w", encoding="utf-8") as f:
        f.write(json_util.dumps(report, indent=2))
    # markdown summary
    lines: List[str] = []
    lines.append("## MongoDB Health Check Report")
    lines.append("")
    lines.append(f"- **generatedAt**: {report['generatedAt']}")
    bi = report.get("server", {}).get("buildInfo", {})
    lines.append(f"- **version**: {bi.get('version', 'unknown')}")
    lines.append("")
    for entry in report.get("collections", []):
        lines.append(f"### {entry['db']}.{entry['collection']}")
        cs = entry.get("collStats", {})
        li = entry.get("largestIndex")
        lines.append(f"- documents: {entry.get('counts', {}).get('documents', 'n/a')}")
        lines.append(f"- storageSize: {cs.get('storageSize', 'n/a')} bytes")
        lines.append(f"- totalIndexSize: {cs.get('totalIndexSize', 'n/a')} bytes")
        if li:
            lines.append(f"- largestIndex: {li['name']} = {li['size']} bytes")
        ds = entry.get("schema", {}).get("docSizeBytes", {})
        lines.append(f"- docSize(max/p95/avg): {ds.get('max', 0)}/{ds.get('p95', 0)}/{int(ds.get('avg', 0))} bytes")
        flags = entry.get("flags", {})
        if flags.get("unboundedArrays"):
            lines.append(f"- WARN unbounded arrays: {', '.join(flags['unboundedArrays'])}")
        if flags.get("largeDocuments"):
            lines.append("- WARN large documents detected (>=12MB)")
        if flags.get("extremeNesting"):
            lines.append("- WARN extreme nesting depth (>10)")
        if entry.get("notes"):
            lines.append(f"- notes: {' | '.join(entry['notes'])}")
        lines.append("")
    if report.get("uriRecommendations"):
        lines.append("## URI Recommendations")
        for r in report["uriRecommendations"]:
            lines.append(f"- {r}")
    # Group by DB and render tables for readability
    groups: Dict[str, List[Dict[str, Any]]] = {}
    for e in report.get("collections", []):
        groups.setdefault(e["db"], []).append(e)

    md_lines: List[str] = []
    md_lines.append("## MongoDB Health Check Report")
    md_lines.append("")
    md_lines.append(f"- **generatedAt**: {report['generatedAt']}")
    bi = report.get("server", {}).get("buildInfo", {})
    md_lines.append(f"- **version**: {bi.get('version', 'unknown')}")
    md_lines.append("")

    for db_name in sorted(groups.keys()):
        md_lines.append(f"## DB: {db_name}")
        md_lines.append("")
        md_lines.append("| Collection | Docs | Storage (B) | TotalIndex (B) | LargestIndex | DocSize max/p95/avg (B) | Warnings |")
        md_lines.append("|---|---:|---:|---:|---|---|---|")
        for entry in sorted(groups[db_name], key=lambda x: x.get("collection", "")):
            cs = entry.get("collStats", {})
            li = entry.get("largestIndex")
            docs = entry.get("counts", {}).get("documents", "n/a")
            storage = cs.get("storageSize", "n/a")
            tix = cs.get("totalIndexSize", "n/a")
            li_str = f"{li['name']} ({li['size']})" if li else "-"
            ds = entry.get("schema", {}).get("docSizeBytes", {})
            ds_str = f"{ds.get('max', 0)}/{ds.get('p95', 0)}/{int(ds.get('avg', 0))}"
            flags = entry.get("flags", {})
            warns: List[str] = []
            if flags.get("unboundedArrays"):
                warns.append("Unbounded arrays: " + ",".join(flags["unboundedArrays"]))
            if flags.get("largeDocuments"):
                warns.append("Large documents")
            if flags.get("extremeNesting"):
                warns.append("Extreme nesting")
            warn_str = "; ".join(warns) if warns else "-"
            md_lines.append(
                f"| {entry['collection']} | {docs} | {storage} | {tix} | {li_str} | {ds_str} | {warn_str} |"
            )
        md_lines.append("")
        # detailed recommendations per collection (tables)
        refs = {
            "cnblogs": "https://www.cnblogs.com/fatedeity/p/16977740.html",
            "mongoing": "https://mongoing.com/archives/82907",
            "manual": "https://www.mongodb.com/zh-cn/docs/manual/data-modeling/schema-design-process/",
        }
        for entry in sorted(groups[db_name], key=lambda x: x.get("collection", "")):
            md_lines.append(f"### {entry['collection']} · Schema Recommendations")
            md_lines.append("")
            md_lines.append("| Issue | Evidence | Recommendation |")
            md_lines.append("|---|---|---|")
            flags = entry.get("flags", {})
            schema = entry.get("schema", {})
            ds = schema.get("docSizeBytes", {})
            depth = schema.get("nestingDepth", {})
            poly = schema.get("polymorphism", {})
            arr = schema.get("arrayStats", {})
            # large documents
            if flags.get("largeDocuments"):
                md_lines.append(f"| Large document size | max={ds.get('max',0)}B, p95={ds.get('p95',0)}B | Split/Reference/Bucket patterns |")
            # extreme nesting
            if flags.get("extremeNesting"):
                md_lines.append(f"| Excessive nesting depth | maxDepth={depth.get('max',0)} | Flatten structure, embed hot fields only |")
            # unbounded arrays
            if flags.get("unboundedArrays"):
                fields_list = ", ".join(flags.get("unboundedArrays", []))
                # pick first field for evidence
                for fpath in flags.get("unboundedArrays", []):
                    st = arr.get(fpath, {})
                    md_lines.append(f"| Potential unbounded array | {fpath}: max={st.get('max',0)}, p95={st.get('p95',0)} | Bucket/Subset/Outlier patterns |")
            # polymorphism
            if poly:
                top_fields = ", ".join(list(poly.keys())[:5])
                md_lines.append(f"| Field type polymorphism | examples: {top_fields} | Normalize types or split by variant |")
            # default when no issues
            if not (flags.get("largeDocuments") or flags.get("extremeNesting") or flags.get("unboundedArrays") or poly):
                md_lines.append("| Healthy | No major schema risks observed in sample | Keep monitoring | - |")
            md_lines.append("")

            # index recommendations
            md_lines.append(f"### {entry['collection']} · Index Recommendations")
            md_lines.append("")
            md_lines.append("| Category | Fields | Note |")
            md_lines.append("|---|---|---|")
            insights = entry.get("indexInsights", {})
            hi = insights.get("highCardinalityCandidates", [])
            lo = insights.get("lowCardinalityWarnings", [])
            if hi:
                md_lines.append(f"| High-cardinality candidates | {', '.join(hi)} | Consider selective leading index keys (if used in filters) |")
            if lo:
                md_lines.append(f"| Low-cardinality warnings | {', '.join(lo)} | Avoid as leading index position; place later in compound |")
            # multikey caution from arrays
            arr_warns = []
            for fpath, st in arr.items():
                if st.get("max", 0) >= 1000 or st.get("p95", 0) >= 500:
                    arr_warns.append(fpath)
            if arr_warns:
                md_lines.append(f"| Multikey caution | {', '.join(sorted(arr_warns))} | Large arrays can cause index entry explosion |")
            if not (hi or lo or arr_warns):
                md_lines.append("| General | - | Validate with $indexStats and plans; maintain lean indexes | - |")
            md_lines.append("")
    if report.get("uriRecommendations"):
        md_lines.append("## URI Recommendations")
        # If serverStatus opcounters are available, present evidence table first
        ss = report.get("server", {}).get("serverStatus", {}) if isinstance(report.get("server"), dict) else {}
        opc = ss.get("opcounters", {}) if isinstance(ss, dict) else {}
        if opc:
            read_count = float(opc.get("query", 0)) + float(opc.get("getmore", 0))
            write_count = float(opc.get("insert", 0)) + float(opc.get("update", 0)) + float(opc.get("delete", 0))
            ratio_val = read_count / write_count if write_count > 0 else (read_count if read_count > 0 else 0.0)
            md_lines.append("")
            md_lines.append("| Metric | Value | Note |")
            md_lines.append("|---|---:|---|")
            md_lines.append(f"| query + getmore | {int(read_count)} | total reads observed (opcounters) |")
            md_lines.append(f"| insert + update + delete | {int(write_count)} | total writes observed (opcounters) |")
            md_lines.append(f"| read/write ratio | {ratio_val:.2f} | >=5 read-heavy; <=0.2 write-heavy |")
            md_lines.append("")
        for r in report["uriRecommendations"]:
            md_lines.append(f"- {r}")
    # references intentionally omitted per user request
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))

    # Also emit bilingual Markdown variants for external sharing
    try:
        en_md = render_markdown(report, lang="en")
        tw_md = render_markdown(report, lang="zh-TW")
        with open(os.path.join(output_dir, "report_en.md"), "w", encoding="utf-8") as f:
            f.write(en_md)
        with open(os.path.join(output_dir, "report_zh-TW.md"), "w", encoding="utf-8") as f:
            f.write(tw_md)
    except Exception:
        pass

    # DOCX exports (English and Traditional Chinese)
    if DOCX_AVAILABLE:
        try:
            export_docx(report, os.path.join(output_dir, "report_en.docx"), lang="en")
            export_docx(report, os.path.join(output_dir, "report_zh-TW.docx"), lang="zh-TW")
        except Exception:
            # do not fail whole run due to docx export
            pass


def export_docx(report: Dict[str, Any], path: str, lang: str = "en") -> None:
    doc = Document()

    def add_title(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_h2(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)

    def add_h3(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)

    # localized strings
    L = {
        "en": {
            "title": "MongoDB Health Check Report",
            "generatedAt": "generatedAt",
            "version": "version",
            "db": "DB",
            "summaryTable": ["Collection", "Docs", "Storage (B)", "TotalIndex (B)", "LargestIndex", "DocSize max/p95/avg (B)", "Warnings"],
            "unbounded": "Unbounded arrays",
            "largeDocs": "Large documents",
            "extremeNest": "Extreme nesting",
            "schemaRec": "Schema Recommendations",
            "schemaCols": ["Issue", "Evidence", "Recommendation"],
            "indexRec": "Index Recommendations",
            "indexCols": ["Category", "Fields", "Note"],
            "healthy": "Healthy",
            "none": "-",
            "uri": "URI Recommendations",
            "metricCols": ["Metric", "Value", "Note"],
        },
        "zh-TW": {
            "title": "MongoDB 健康檢查報告",
            "generatedAt": "產生時間",
            "version": "版本",
            "db": "資料庫",
            "summaryTable": ["集合", "文件數", "儲存位元組", "索引總大小", "最大索引", "文件大小 最大/95分位/平均", "警示"],
            "unbounded": "無上界陣列",
            "largeDocs": "大型文件",
            "extremeNest": "過深巢狀",
            "schemaRec": "Schema 建議",
            "schemaCols": ["問題", "證據", "建議"],
            "indexRec": "索引建議",
            "indexCols": ["類別", "欄位", "說明"],
            "healthy": "健康",
            "none": "-",
            "uri": "URI 連線建議",
            "metricCols": ["指標", "數值", "說明"],
        },
    }[lang]

    add_title(L["title"])
    p = doc.add_paragraph()
    p.add_run(f"- {L['generatedAt']}: {report.get('generatedAt','')}")
    bi = report.get("server", {}).get("buildInfo", {})
    p.add_run(f"\n- {L['version']}: {bi.get('version','unknown')}")

    # group collections by DB
    groups: Dict[str, List[Dict[str, Any]]] = {}
    for e in report.get("collections", []):
        groups.setdefault(e["db"], []).append(e)

    for db_name in sorted(groups.keys()):
        add_h2(f"{L['db']}: {db_name}")
        # summary table
        tbl = doc.add_table(rows=1, cols=len(L["summaryTable"]))
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(L["summaryTable"]):
            hdr_cells[i].text = h
        for entry in sorted(groups[db_name], key=lambda x: x.get("collection", "")):
            cs = entry.get("collStats", {})
            li = entry.get("largestIndex")
            docs = entry.get("counts", {}).get("documents", "n/a")
            storage = cs.get("storageSize", "n/a")
            tix = cs.get("totalIndexSize", "n/a")
            li_str = f"{li['name']} ({li['size']})" if li else L["none"]
            ds = entry.get("schema", {}).get("docSizeBytes", {})
            ds_str = f"{ds.get('max', 0)}/{ds.get('p95', 0)}/{int(ds.get('avg', 0))}"
            flags = entry.get("flags", {})
            warns: List[str] = []
            if flags.get("unboundedArrays"):
                warns.append(f"{L['unbounded']}: {', '.join(flags['unboundedArrays'])}")
            if flags.get("largeDocuments"):
                warns.append(L["largeDocs"])
            if flags.get("extremeNesting"):
                warns.append(L["extremeNest"])
            warn_str = "; ".join(warns) if warns else L["none"]
            row = tbl.add_row().cells
            row[0].text = entry["collection"]
            row[1].text = str(docs)
            row[2].text = str(storage)
            row[3].text = str(tix)
            row[4].text = li_str
            row[5].text = ds_str
            row[6].text = warn_str

        # per-collection recommendations
        for entry in sorted(groups[db_name], key=lambda x: x.get("collection", "")):
            add_h3(f"{entry['collection']} · {L['schemaRec']}")
            tbls = doc.add_table(rows=1, cols=3)
            hdr = tbls.rows[0].cells
            for i, h in enumerate(L["schemaCols"]):
                hdr[i].text = h
            flags = entry.get("flags", {})
            schema = entry.get("schema", {})
            ds = schema.get("docSizeBytes", {})
            depth = schema.get("nestingDepth", {})
            poly = schema.get("polymorphism", {})
            arr = schema.get("arrayStats", {})
            # large docs
            if flags.get("largeDocuments"):
                r = tbls.add_row().cells
                r[0].text = {
                    "en": "Large document size",
                    "zh-TW": "大型文件尺寸",
                }[lang]
                r[1].text = f"max={ds.get('max',0)}B, p95={ds.get('p95',0)}B"
                r[2].text = {
                    "en": "Split/Reference/Bucket patterns",
                    "zh-TW": "拆分/引用/分桶 模式",
                }[lang]
            # extreme nesting
            if flags.get("extremeNesting"):
                r = tbls.add_row().cells
                r[0].text = {
                    "en": "Excessive nesting depth",
                    "zh-TW": "過深巢狀層級",
                }[lang]
                r[1].text = f"maxDepth={depth.get('max',0)}"
                r[2].text = {
                    "en": "Flatten structure; embed hot fields only",
                    "zh-TW": "扁平化結構；僅內嵌高頻欄位",
                }[lang]
            # unbounded arrays
            if flags.get("unboundedArrays"):
                for fpath in flags.get("unboundedArrays", []):
                    st = arr.get(fpath, {})
                    r = tbls.add_row().cells
                    r[0].text = {
                        "en": "Potential unbounded array",
                        "zh-TW": "疑似無上界陣列",
                    }[lang]
                    r[1].text = f"{fpath}: max={st.get('max',0)}, p95={st.get('p95',0)}"
                    r[2].text = {
                        "en": "Bucket/Subset/Outlier patterns",
                        "zh-TW": "分桶/子集/異常外移 模式",
                    }[lang]
            # polymorphism
            if poly:
                r = tbls.add_row().cells
                r[0].text = {
                    "en": "Field type polymorphism",
                    "zh-TW": "欄位型別多態",
                }[lang]
                r[1].text = ", ".join(list(poly.keys())[:5])
                r[2].text = {
                    "en": "Normalize types or split by variant",
                    "zh-TW": "統一型別或依變體拆分",
                }[lang]
            if not (flags.get("largeDocuments") or flags.get("extremeNesting") or flags.get("unboundedArrays") or poly):
                r = tbls.add_row().cells
                r[0].text = L["healthy"]
                r[1].text = {
                    "en": "No major schema risks observed in sample",
                    "zh-TW": "樣本未觀察到主要風險",
                }[lang]
                r[2].text = {
                    "en": "Keep monitoring",
                    "zh-TW": "持續監控",
                }[lang]

            # index recommendations
            add_h3(f"{entry['collection']} · {L['indexRec']}")
            tbli = doc.add_table(rows=1, cols=3)
            ih = tbli.rows[0].cells
            for i, h in enumerate(L["indexCols"]):
                ih[i].text = h
            insights = entry.get("indexInsights", {})
            hi = insights.get("highCardinalityCandidates", [])
            lo = insights.get("lowCardinalityWarnings", [])
            arr_warns: List[str] = []
            arr = schema.get("arrayStats", {})
            for fpath, st in arr.items():
                if st.get("max", 0) >= 1000 or st.get("p95", 0) >= 500:
                    arr_warns.append(fpath)
            if hi:
                r = tbli.add_row().cells
                r[0].text = {
                    "en": "High-cardinality candidates",
                    "zh-TW": "高基數候選",
                }[lang]
                r[1].text = ", ".join(hi)
                r[2].text = {
                    "en": "Consider selective leading index keys (if used in filters)",
                    "zh-TW": "若用於篩選，建議作為索引前導鍵",
                }[lang]
            if lo:
                r = tbli.add_row().cells
                r[0].text = {
                    "en": "Low-cardinality warnings",
                    "zh-TW": "低基數警示",
                }[lang]
                r[1].text = ", ".join(lo)
                r[2].text = {
                    "en": "Avoid as leading index position; place later in compound",
                    "zh-TW": "避免作為複合索引前導鍵，置於後位",
                }[lang]
            if arr_warns:
                r = tbli.add_row().cells
                r[0].text = {
                    "en": "Multikey caution",
                    "zh-TW": "多鍵索引警示",
                }[lang]
                r[1].text = ", ".join(sorted(arr_warns))
                r[2].text = {
                    "en": "Large arrays can cause index entry explosion",
                    "zh-TW": "巨大陣列可能導致索引條目爆炸",
                }[lang]
            if not (hi or lo or arr_warns):
                r = tbli.add_row().cells
                r[0].text = {
                    "en": "General",
                    "zh-TW": "一般",
                }[lang]
                r[1].text = L["none"]
                r[2].text = {
                    "en": "Validate with $indexStats and plans; maintain lean indexes",
                    "zh-TW": "透過 $indexStats 與執行計劃驗證；維持精簡索引",
                }[lang]

    # URI recommendations section
    if report.get("uriRecommendations"):
        add_h2(L["uri"])
        ss = report.get("server", {}).get("serverStatus", {}) if isinstance(report.get("server"), dict) else {}
        opc = ss.get("opcounters", {}) if isinstance(ss, dict) else {}
        if opc:
            tblm = doc.add_table(rows=1, cols=3)
            mh = tblm.rows[0].cells
            for i, h in enumerate(L["metricCols"]):
                mh[i].text = h
            read_count = float(opc.get("query", 0)) + float(opc.get("getmore", 0))
            write_count = float(opc.get("insert", 0)) + float(opc.get("update", 0)) + float(opc.get("delete", 0))
            ratio_val = read_count / write_count if write_count > 0 else (read_count if read_count > 0 else 0.0)
            r = tblm.add_row().cells
            r[0].text = {"en": "query + getmore", "zh-TW": "查詢+取回"}[lang]
            r[1].text = str(int(read_count))
            r[2].text = {"en": "total reads observed (opcounters)", "zh-TW": "觀測讀取總數（opcounters）"}[lang]
            r = tblm.add_row().cells
            r[0].text = {"en": "insert + update + delete", "zh-TW": "插入+更新+刪除"}[lang]
            r[1].text = str(int(write_count))
            r[2].text = {"en": "total writes observed (opcounters)", "zh-TW": "觀測寫入總數（opcounters）"}[lang]
            r = tblm.add_row().cells
            r[0].text = {"en": "read/write ratio", "zh-TW": "讀寫比"}[lang]
            r[1].text = f"{ratio_val:.2f}"
            r[2].text = {"en": ">=5 read-heavy; <=0.2 write-heavy", "zh-TW": ">=5 讀多；<=0.2 寫多"}[lang]
        # bullet list of uri recommendations
        for rec in report.get("uriRecommendations", []):
            doc.add_paragraph(f"- {rec}")

    doc.save(path)


def render_markdown(report: Dict[str, Any], lang: str = "en") -> str:
    L = {
        "en": {
            "title": "MongoDB Health Check Report",
            "generatedAt": "generatedAt",
            "version": "version",
            "db": "DB",
            "summaryHeader": "| Collection | Docs | Storage (B) | TotalIndex (B) | LargestIndex | DocSize max/p95/avg (B) | Warnings |\n|---|---:|---:|---:|---|---|---|",
            "unbounded": "Unbounded arrays",
            "largeDocs": "Large documents",
            "extremeNest": "Extreme nesting",
            "schemaRec": "Schema Recommendations",
            "schemaCols": "| Issue | Evidence | Recommendation |\n|---|---|---|",
            "indexRec": "Index Recommendations",
            "indexCols": "| Category | Fields | Note |\n|---|---|---|",
            "healthy": "Healthy",
            "healthyEv": "No major schema risks observed in sample",
            "healthyRec": "Keep monitoring",
            "general": "General",
            "uri": "URI Recommendations",
            "metricCols": "| Metric | Value | Note |\n|---|---:|---|",
            "reads": "query + getmore",
            "writes": "insert + update + delete",
            "rw": "read/write ratio",
            "readsNote": "total reads observed (opcounters)",
            "writesNote": "total writes observed (opcounters)",
            "rwNote": ">=5 read-heavy; <=0.2 write-heavy",
        },
        "zh-TW": {
            "title": "MongoDB 健康檢查報告",
            "generatedAt": "產生時間",
            "version": "版本",
            "db": "資料庫",
            "summaryHeader": "| 集合 | 文件數 | 儲存位元組 | 索引總大小 | 最大索引 | 文件大小 最大/95分位/平均 | 警示 |\n|---|---:|---:|---:|---|---|---|",
            "unbounded": "無上界陣列",
            "largeDocs": "大型文件",
            "extremeNest": "過深巢狀",
            "schemaRec": "Schema 建議",
            "schemaCols": "| 問題 | 證據 | 建議 |\n|---|---|---|",
            "indexRec": "索引建議",
            "indexCols": "| 類別 | 欄位 | 說明 |\n|---|---|---|",
            "healthy": "健康",
            "healthyEv": "樣本未觀察到主要風險",
            "healthyRec": "持續監控",
            "general": "一般",
            "uri": "URI 連線建議",
            "metricCols": "| 指標 | 數值 | 說明 |\n|---|---:|---|",
            "reads": "查詢 + 取回",
            "writes": "插入 + 更新 + 刪除",
            "rw": "讀寫比",
            "readsNote": "觀測讀取總數（opcounters）",
            "writesNote": "觀測寫入總數（opcounters）",
            "rwNote": ">=5 讀多；<=0.2 寫多",
        },
    }[lang]

    groups: Dict[str, List[Dict[str, Any]]] = {}
    for e in report.get("collections", []):
        groups.setdefault(e["db"], []).append(e)

    out: List[str] = []
    out.append(f"## {L['title']}")
    out.append("")
    out.append(f"- **{L['generatedAt']}**: {report.get('generatedAt','')}")
    bi = report.get("server", {}).get("buildInfo", {})
    out.append(f"- **{L['version']}**: {bi.get('version','unknown')}")
    out.append("")
    for db_name in sorted(groups.keys()):
        out.append(f"## {L['db']}: {db_name}")
        out.append("")
        out.append(L["summaryHeader"])
        for entry in sorted(groups[db_name], key=lambda x: x.get("collection", "")):
            cs = entry.get("collStats", {})
            li = entry.get("largestIndex")
            docs = entry.get("counts", {}).get("documents", "n/a")
            storage = cs.get("storageSize", "n/a")
            tix = cs.get("totalIndexSize", "n/a")
            li_str = f"{li['name']} ({li['size']})" if li else "-"
            ds = entry.get("schema", {}).get("docSizeBytes", {})
            ds_str = f"{ds.get('max', 0)}/{ds.get('p95', 0)}/{int(ds.get('avg', 0))}"
            flags = entry.get("flags", {})
            warns: List[str] = []
            if flags.get("unboundedArrays"):
                warns.append(f"{L['unbounded']}: " + ",".join(flags["unboundedArrays"]))
            if flags.get("largeDocuments"):
                warns.append(L["largeDocs"])
            if flags.get("extremeNesting"):
                warns.append(L["extremeNest"])
            warn_str = "; ".join(warns) if warns else "-"
            out.append(f"| {entry['collection']} | {docs} | {storage} | {tix} | {li_str} | {ds_str} | {warn_str} |")
        out.append("")
        for entry in sorted(groups[db_name], key=lambda x: x.get("collection", "")):
            out.append(f"### {entry['collection']} · {L['schemaRec']}")
            out.append("")
            out.append(L["schemaCols"])
            flags = entry.get("flags", {})
            schema = entry.get("schema", {})
            ds = schema.get("docSizeBytes", {})
            depth = schema.get("nestingDepth", {})
            poly = schema.get("polymorphism", {})
            arr = schema.get("arrayStats", {})
            if flags.get("largeDocuments"):
                out.append(f"| Large document size | max={ds.get('max',0)}B, p95={ds.get('p95',0)}B | Split/Reference/Bucket patterns |")
            if flags.get("extremeNesting"):
                out.append(f"| Excessive nesting depth | maxDepth={depth.get('max',0)} | Flatten structure, embed hot fields only |")
            if flags.get("unboundedArrays"):
                for fpath in flags.get("unboundedArrays", []):
                    st = arr.get(fpath, {})
                    out.append(f"| Potential unbounded array | {fpath}: max={st.get('max',0)}, p95={st.get('p95',0)} | Bucket/Subset/Outlier patterns |")
            if poly:
                top_fields = ", ".join(list(poly.keys())[:5])
                out.append(f"| Field type polymorphism | examples: {top_fields} | Normalize types or split by variant |")
            if not (flags.get("largeDocuments") or flags.get("extremeNesting") or flags.get("unboundedArrays") or poly):
                out.append(f"| {L['healthy']} | {L['healthyEv']} | {L['healthyRec']} |")
            out.append("")
            out.append(f"### {entry['collection']} · {L['indexRec']}")
            out.append("")
            out.append(L["indexCols"])
            insights = entry.get("indexInsights", {})
            hi = insights.get("highCardinalityCandidates", [])
            lo = insights.get("lowCardinalityWarnings", [])
            arr_warns = []
            for fpath, st in arr.items():
                if st.get("max", 0) >= 1000 or st.get("p95", 0) >= 500:
                    arr_warns.append(fpath)
            if hi:
                out.append(f"| High-cardinality candidates | {', '.join(hi)} | Consider selective leading index keys (if used in filters) |")
            if lo:
                out.append(f"| Low-cardinality warnings | {', '.join(lo)} | Avoid as leading index position; place later in compound |")
            if arr_warns:
                out.append(f"| Multikey caution | {', '.join(sorted(arr_warns))} | Large arrays can cause index entry explosion |")
            if not (hi or lo or arr_warns):
                out.append(f"| {L['general']} | - | Validate with $indexStats and plans; maintain lean indexes |")
            out.append("")
    if report.get("uriRecommendations"):
        out.append(f"## {L['uri']}")
        ss = report.get("server", {}).get("serverStatus", {}) if isinstance(report.get("server"), dict) else {}
        opc = ss.get("opcounters", {}) if isinstance(ss, dict) else {}
        if opc:
            read_count = float(opc.get("query", 0)) + float(opc.get("getmore", 0))
            write_count = float(opc.get("insert", 0)) + float(opc.get("update", 0)) + float(opc.get("delete", 0))
            ratio_val = read_count / write_count if write_count > 0 else (read_count if read_count > 0 else 0.0)
            out.append("")
            out.append(L["metricCols"])
            out.append(f"| {L['reads']} | {int(read_count)} | {L['readsNote']} |")
            out.append(f"| {L['writes']} | {int(write_count)} | {L['writesNote']} |")
            out.append(f"| {L['rw']} | {ratio_val:.2f} | {L['rwNote']} |")
            out.append("")
        for r in report["uriRecommendations"]:
            out.append(f"- {r}")
    return "\n".join(out)


def main() -> int:
    args = parse_args()
    # Optional .env loading
    env_path = args.env_file or os.environ.get("ENV_FILE", "")
    if env_path:
        try:
            from dotenv import load_dotenv  # type: ignore
            load_dotenv(dotenv_path=env_path, override=False)
        except Exception:
            pass
    else:
        # load default .env if exists in project root
        try:
            from dotenv import load_dotenv  # type: ignore
            load_dotenv(override=False)
        except Exception:
            pass

    # Determine MongoDB URI: CLI > env var
    effective_uri = args.uri.strip() or os.environ.get("MONGODB_URI", "").strip()
    if not effective_uri:
        print("Missing MongoDB URI. Provide --uri or set MONGODB_URI in environment or .env", file=sys.stderr)
        return 2
    rng = random.Random(args.seed)
    client = MongoClient(effective_uri, serverSelectionTimeoutMS=args.timeout_ms)
    report: Dict[str, Any] = {
        "generatedAt": datetime.datetime.now(datetime.timezone.utc).isoformat().replace("+00:00", "Z"),
        "server": {},
        "collections": [],
        "uriRecommendations": [],
        "errors": [],
    }
    try:
        client.admin.command("ping")
    except PyMongoError as exc:
        print(f"Connection failed: {exc}", file=sys.stderr)
        report["errors"].append(f"Connection failed: {exc}")
        write_reports(args.output_dir, report)
        return 2

    # server info
    server_info = get_server_info(client)
    report["server"] = server_info
    # recommendations
    report["uriRecommendations"] = recommend_uri_tweaks(effective_uri, server_info)

    # targets
    include_dbs = [s for s in (x.strip() for x in args.dbs.split(",")) if s]
    include_colls = [s for s in (x.strip() for x in args.collections.split(",")) if s]

    # build document-level filter if requested
    query_filter: Optional[Dict[str, Any]] = None
    if args.doc_id:
        # try ObjectId, else use as string
        try:
            oid = ObjectId(args.doc_id)
            query_filter = {"_id": oid}
        except Exception:
            query_filter = {"_id": args.doc_id}
    elif args.filter:
        try:
            query_filter = json.loads(args.filter)
        except Exception as exc:
            report["errors"].append(f"Invalid --filter JSON: {exc}")
            query_filter = None
    targets = get_target_sets(client, include_dbs, include_colls)
    if not targets:
        # fallback: try a default db hinted by URI (e.g., 'test' on Atlas when none specified is accessible via list)
        try:
            default_db_name = client.get_default_database().name  # type: ignore[attr-defined]
            targets = {default_db_name: client[default_db_name].list_collection_names()}
        except Exception:
            pass

    for db_name, colls in targets.items():
        db = client[db_name]
        for coll_name in colls:
            coll = db[coll_name]
            try:
                entry = analyze_collection(db_name, coll_name, coll, args.sample_size, args.max_docs_per_coll, rng, query_filter=query_filter)
                report["collections"].append(entry)
            except Exception as exc:
                report["errors"].append(f"Analyze failed for {db_name}.{coll_name}: {exc}")

    write_reports(args.output_dir, report)
    print(f"Report written to {args.output_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())


